"""DOCX loader — extracts paragraph text and merges sidecar metadata."""

from __future__ import annotations

from pathlib import Path

import docx

from .frontmatter import FrontMatterError
from .markdown_loader import _coerce_str, _derive_document_id, _derive_policy_family
from .models import AccountingDocument, compute_checksum
from .sidecar import load_sidecar_metadata


_REQUIRED_FIELDS = ("title", "version", "document_type")


def _extract_text(path: Path) -> str:
    document = docx.Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts).strip()


def load_docx_document(path: Path) -> AccountingDocument:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"file not found: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"not a DOCX file: {path}")

    metadata = load_sidecar_metadata(path)
    missing = [f for f in _REQUIRED_FIELDS if not metadata.get(f)]
    if missing:
        raise FrontMatterError(
            f"{path}: sidecar metadata missing required field(s): {missing}"
        )

    content = _extract_text(path)
    if not content:
        raise FrontMatterError(
            f"{path}: DOCX paragraphs are empty. Empty documents are not "
            "supported (the chunker needs non-empty content)."
        )

    document_id = _derive_document_id(metadata, path)
    policy_family = _derive_policy_family(metadata, path)
    is_malicious = bool(metadata.get("is_malicious", False))

    canonical_metadata = {
        "document_id": document_id,
        "title": metadata.get("title"),
        "version": metadata.get("version"),
        "valid_from": metadata.get("valid_from"),
        "valid_to": metadata.get("valid_to"),
        "document_type": metadata.get("document_type"),
        "client": metadata.get("client"),
        "policy_family": policy_family,
        "replaces": metadata.get("replaces"),
        "risk_type": metadata.get("risk_type"),
        "is_malicious": is_malicious,
    }
    checksum = compute_checksum(content, canonical_metadata)

    extra_metadata = {k: v for k, v in metadata.items() if k not in canonical_metadata}
    extra_metadata["source_format"] = "docx"

    return AccountingDocument(
        document_id=document_id,
        title=str(metadata.get("title")),
        version=str(metadata.get("version")),
        valid_from=_coerce_str(metadata.get("valid_from")),
        valid_to=_coerce_str(metadata.get("valid_to")),
        document_type=str(metadata.get("document_type")),
        client=_coerce_str(metadata.get("client")),
        policy_family=policy_family,
        replaces=_coerce_str(metadata.get("replaces")),
        risk_type=_coerce_str(metadata.get("risk_type")),
        is_malicious=is_malicious,
        source_path=str(path),
        content=content,
        checksum=checksum,
        metadata=extra_metadata,
    )
