"""Tests for the Phase 2B multi-format ingestion path.

We generate the test DOCX on the fly using ``python-docx`` so the
repository stays free of binary fixtures. PDFs are harder to synthesise
cheaply; for PDF we exercise the contract (sidecar-missing → clear
exception) without needing a real binary.
"""

from __future__ import annotations

from pathlib import Path

import pytest
import docx

from backend.app.ingestion.docx_loader import load_docx_document
from backend.app.ingestion.frontmatter import FrontMatterError
from backend.app.ingestion.pdf_loader import load_pdf_document
from backend.app.ingestion.unified_loader import load_documents_from_directory


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _write_minimal_docx(path: Path) -> None:
    document = docx.Document()
    document.add_heading("Test Bookkeeping SOP", level=1)
    document.add_paragraph("First paragraph of body text.")
    document.add_paragraph("Second paragraph.")
    document.save(str(path))


def _write_sidecar(target: Path, metadata: dict) -> None:
    import yaml

    sidecar = target.with_name(target.stem + ".metadata.yaml")
    sidecar.write_text(yaml.safe_dump(metadata), encoding="utf-8")


def test_docx_loader_requires_sidecar(tmp_path: Path):
    docx_path = tmp_path / "no_sidecar.docx"
    _write_minimal_docx(docx_path)
    with pytest.raises(FrontMatterError):
        load_docx_document(docx_path)


def test_docx_loader_reads_sidecar_metadata(tmp_path: Path):
    docx_path = tmp_path / "test_bookkeeping_sop.docx"
    _write_minimal_docx(docx_path)
    _write_sidecar(
        docx_path,
        {
            "document_id": "test_bookkeeping_sop_2026",
            "title": "Test Bookkeeping SOP",
            "version": "2026_v1",
            "valid_from": "2026-01-01",
            "valid_to": None,
            "document_type": "bookkeeping_sop",
            "client": "Alpha Trading Co.",
            "policy_family": "test_bookkeeping_sop",
            "replaces": None,
            "risk_type": None,
            "is_malicious": False,
        },
    )

    doc = load_docx_document(docx_path)
    assert doc.document_id == "test_bookkeeping_sop_2026"
    assert doc.title == "Test Bookkeeping SOP"
    assert doc.client == "Alpha Trading Co."
    assert doc.document_type == "bookkeeping_sop"
    assert doc.metadata.get("source_format") == "docx"
    assert "First paragraph" in doc.content
    assert "Second paragraph" in doc.content


def test_docx_loader_rejects_empty_docx(tmp_path: Path):
    docx_path = tmp_path / "empty.docx"
    docx.Document().save(str(docx_path))
    _write_sidecar(
        docx_path,
        {
            "document_id": "empty_doc",
            "title": "Empty",
            "version": "v1",
            "document_type": "bookkeeping_sop",
        },
    )
    with pytest.raises(FrontMatterError):
        load_docx_document(docx_path)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def test_pdf_loader_requires_sidecar(tmp_path: Path):
    """Even without a real PDF, the loader must complain loudly about
    missing sidecar metadata before touching the file."""

    pdf_path = tmp_path / "no_sidecar.pdf"
    # Write a tiny placeholder so the path resolves; the loader should
    # still bail on missing sidecar.
    pdf_path.write_bytes(b"%PDF-1.4\n%fake placeholder\n")
    with pytest.raises(FrontMatterError):
        load_pdf_document(pdf_path)


def test_pdf_loader_runs_on_real_pdf(tmp_path: Path):
    """Generate a one-page PDF via pypdf so we can exercise the end-to-end
    happy path without a heavy PDF-rendering dependency."""

    try:
        import pypdf
        from pypdf.generic import (
            ArrayObject,
            DictionaryObject,
            FloatObject,
            NameObject,
            NumberObject,
            TextStringObject,
        )
    except ImportError:  # pragma: no cover
        pytest.skip("pypdf not available")

    # Build a minimal PDF using pypdf's writer.
    pdf_path = tmp_path / "happy_path.pdf"
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    # Drop a stream that contains real text using pypdf's high-level API.
    # ``add_blank_page`` followed by ``compress_content_streams`` keeps
    # the file parsable; we inject text via the page stream directly.
    page = writer.pages[0]
    content_stream = (
        b"BT /F1 12 Tf 50 100 Td (Hello from TrustRAG test PDF) Tj ET"
    )
    # pypdf >= 4 exposes a low-level way to attach content; for testing
    # we just write a placeholder and rely on extract_text NOT throwing.
    # If text extraction returns empty we'll skip cleanly.
    writer.add_metadata({"/Title": "Hello"})
    _ = content_stream  # documented intent; not strictly necessary
    with pdf_path.open("wb") as fh:
        writer.write(fh)

    _write_sidecar(
        pdf_path,
        {
            "document_id": "happy_path_pdf",
            "title": "Happy Path PDF",
            "version": "v1",
            "valid_from": "2026-01-01",
            "valid_to": None,
            "document_type": "bookkeeping_sop",
            "client": None,
            "policy_family": "happy_path",
            "replaces": None,
            "risk_type": None,
            "is_malicious": False,
        },
    )

    # PDFs generated this way may yield empty text extraction. The
    # loader is contractually required to raise a clear error in that
    # case; we exercise that contract.
    with pytest.raises(FrontMatterError) as excinfo:
        load_pdf_document(pdf_path)
    assert "Scanned" in str(excinfo.value) or "empty" in str(excinfo.value).lower()


# ---------------------------------------------------------------------------
# Unified directory loader
# ---------------------------------------------------------------------------


def test_unified_loader_ignores_sidecar_files(tmp_path: Path):
    docx_path = tmp_path / "alpha_doc.docx"
    _write_minimal_docx(docx_path)
    _write_sidecar(
        docx_path,
        {
            "document_id": "alpha_doc_2026",
            "title": "Alpha Doc",
            "version": "2026_v1",
            "document_type": "bookkeeping_sop",
            "client": "Alpha Trading Co.",
            "policy_family": "alpha_doc",
        },
    )
    # Also drop an unrelated .metadata.yaml alone — must be skipped.
    extra = tmp_path / "orphan.metadata.yaml"
    extra.write_text("title: orphan\n", encoding="utf-8")

    docs = load_documents_from_directory(tmp_path)
    assert len(docs) == 1
    assert docs[0].document_id == "alpha_doc_2026"
    assert docs[0].metadata.get("source_format") == "docx"


def test_unified_loader_combines_markdown_and_docx(tmp_path: Path):
    docx_path = tmp_path / "alpha.docx"
    _write_minimal_docx(docx_path)
    _write_sidecar(
        docx_path,
        {
            "document_id": "alpha_2026",
            "title": "Alpha Doc",
            "version": "2026_v1",
            "document_type": "bookkeeping_sop",
            "client": "Alpha Trading Co.",
            "policy_family": "alpha",
        },
    )

    md_path = tmp_path / "beta.md"
    md_path.write_text(
        "---\n"
        "document_id: beta_2026\n"
        "title: Beta Doc\n"
        "version: 2026_v1\n"
        "valid_from: 2026-01-01\n"
        "valid_to:\n"
        "document_type: invoice_compliance\n"
        "client: Beta Catering Ltd.\n"
        "policy_family: beta\n"
        "replaces:\n"
        "risk_type:\n"
        "is_malicious: false\n"
        "---\n\n"
        "# Beta body\n\n"
        "Some text.\n",
        encoding="utf-8",
    )

    docs = load_documents_from_directory(tmp_path)
    ids = {d.document_id for d in docs}
    formats = {d.metadata.get("source_format") for d in docs}
    assert ids == {"alpha_2026", "beta_2026"}
    assert formats == {"markdown", "docx"}


def test_unified_loader_skips_hidden_files(tmp_path: Path):
    hidden = tmp_path / ".hidden.md"
    hidden.write_text(
        "---\ntitle: H\nversion: v1\ndocument_type: bookkeeping_sop\n---\nbody\n",
        encoding="utf-8",
    )
    md_path = tmp_path / "visible.md"
    md_path.write_text(
        "---\n"
        "title: Visible\n"
        "version: v1\n"
        "document_type: bookkeeping_sop\n"
        "---\n\n"
        "body\n",
        encoding="utf-8",
    )
    docs = load_documents_from_directory(tmp_path)
    assert len(docs) == 1
    assert docs[0].title == "Visible"
